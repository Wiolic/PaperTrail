#!/usr/bin/env python3
"""
把本库 topics/ 综述风格的 Markdown(标题/段落/粗体/[[citekey]]反链/表格/分隔线)
转成排版好的 .docx, 供导出给不用 Markdown 的读者。

用法:
  python scripts/md_to_docx.py topics/xxx.md --out exports/xxx.docx --title "文档标题(可选,默认取第一个#标题)"

支持的 Markdown 子集(本库 topics/ 综述实际会用到的):
  # / ## / ### 标题
  纯文本段落, 支持行内 **粗体**
  [[citekey]] 反链 -> 渲染成斜体等宽小字, 不依赖笔记是否存在(纯文本渲染, 非超链接)
  | a | b | c | 表格(含表头分隔行 |---|---|)
  - / * 开头的无序列表
  --- 单独一行 == 分隔线(用段落下边框模拟)
不支持: 嵌套列表、脚注、图片、代码块——本库综述目前不用这些, 需要时再扩展。
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

CITEKEY_RE = re.compile(r"\[\[([^\]]+)\]\]")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_runs(paragraph, text: str):
    """把一段可能含 **粗体** 和 [[citekey]] 的文本拆成 run 加入段落。"""
    # 先按 citekey 切, 再对每一段按粗体切
    pos = 0
    for m in CITEKEY_RE.finditer(text):
        _add_bold_runs(paragraph, text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.italic = True
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x44, 0x55, 0x88)
        pos = m.end()
    _add_bold_runs(paragraph, text[pos:])


def _add_bold_runs(paragraph, text: str):
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        b = paragraph.add_run(m.group(1))
        b.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not re.match(r"^:?-+:?$", row[0]):  # 跳过表头分隔行
            rows.append(row)
        i += 1
    return rows, i


def convert(md_path: Path, title: str | None) -> Document:
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    first_h1 = None
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            add_bottom_border(p)
            i += 1
            continue

        if stripped.startswith("| "):
            rows, i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        if c >= len(rows[0]):
                            continue
                        cell = table.cell(r, c)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs(p, cell_text)
                        if r == 0:
                            for run in p.runs:
                                run.bold = True
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            if level == 1 and first_h1 is None:
                first_h1 = heading_text
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run_text = re.sub(r"^>\s?", "", stripped)
            add_runs(p, run_text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        if stripped.startswith("*(") and stripped.endswith(")*"):
            p = doc.add_paragraph()
            add_runs(p, stripped.strip("*"))
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    if title:
        doc.core_properties.title = title
    elif first_h1:
        doc.core_properties.title = first_h1

    return doc


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("md_file")
    ap.add_argument("--out", required=True)
    ap.add_argument("--title", default=None)
    args = ap.parse_args()

    md_path = Path(args.md_file)
    if not md_path.exists():
        sys.exit(f"文件不存在: {md_path}")

    doc = convert(md_path, args.title)
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"已生成 {out_path}")


if __name__ == "__main__":
    main()
